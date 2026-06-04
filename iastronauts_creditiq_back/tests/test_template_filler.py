"""
Tests for the deterministic .docx template filler (Agent 4).

The integration test loads the *real* corporate template and a *real* ScorerOutput
fixture, fills it the same way the handler does (deterministic + narrative
fallbacks, no LLM), and asserts that no `{{TOKEN}}` survives — which is exactly
the failure that produced the "No .docx found" download error.
"""

import io
import json
import os

import pytest
from docx import Document

pytestmark = pytest.mark.filterwarnings("ignore")

from agents.report_generator.handler import _extract_placeholders, _fill_docx_template
from agents.report_generator.template_filler import (
    build_deterministic_fields,
    build_narrative_fallbacks,
    is_narrative_field,
    _mm,
    _qty,
    _pct,
    _period_label,
)

_FIX = os.path.join(os.path.dirname(__file__), "fixtures")
_TEMPLATE = os.path.join(_FIX, "CreditIQ_Template_EEFF.docx")
_SAMPLE = os.path.join(_FIX, "risk_scorer_sample.json")


def _load_payload():
    from shared.models import ScorerOutput

    with open(_SAMPLE, encoding="utf-8") as f:
        return ScorerOutput.model_validate(json.load(f))


def _merge_like_handler(placeholders, det_map, fallback_map, llm_map=None):
    """Replicate handler precedence: LLM > deterministic > fallback > ''."""
    llm_map = llm_map or {}
    field_map = {}
    for ph in placeholders:
        llm_val = llm_map.get(ph)
        if llm_val and llm_val.strip():
            field_map[ph] = llm_val
        elif ph in det_map:
            field_map[ph] = det_map[ph]
        elif ph in fallback_map:
            field_map[ph] = fallback_map[ph]
        else:
            field_map[ph] = ""
    return field_map


# ---------------------------------------------------------------------------
# Unit tests
# ---------------------------------------------------------------------------

def test_is_narrative_field_classification():
    assert is_narrative_field("MACRO_CONTEXT")
    assert is_narrative_field("BS_ASSETS_ANALYSIS")
    assert is_narrative_field("FINDING_3_TITLE")
    assert is_narrative_field("FINDING_3_BODY")
    # Deterministic — must NOT be sent to the LLM.
    assert not is_narrative_field("BS_R1_C1")
    assert not is_narrative_field("FINDING_3_ACCOUNTS")
    assert not is_narrative_field("FINDING_3_IMPACT")
    assert not is_narrative_field("KPI_NET_INCOME")


def test_money_and_pct_formatting():
    # Values are already in COP MM — rendered as-is, not divided.
    assert _mm(2_904_080.3672) == "2,904,080.4"  # COP MM
    assert _mm(15617.878) == "15,617.9"
    assert _mm(None) == "N/D"
    # Adaptive precision keeps small holdings visible instead of rounding to 0.0.
    assert _mm(0.5) == "0.500"
    assert _mm(0.0000156) == "0.0000156"  # ~3 significant figures, never 0.0
    assert _mm(0.0008) == "0.000800"
    assert _mm(0) == "0.0"
    # Nominal quantities render plainly (no MM scaling).
    assert _qty(1_000_000) == "1,000,000"
    assert _qty(None) == "N/D"
    assert _pct(12.34) == "+12.3%"
    assert _pct(-5.0) == "-5.0%"
    assert _pct(None) == "N/D"
    assert _period_label("2025-06") == "Jun 2025"


# ---------------------------------------------------------------------------
# Integration: fill the real template with a real payload, no LLM
# ---------------------------------------------------------------------------

@pytest.mark.skipif(not os.path.exists(_TEMPLATE) or not os.path.exists(_SAMPLE),
                    reason="template/fixture not present")
def test_real_template_fully_filled_without_llm():
    payload = _load_payload()
    with open(_TEMPLATE, "rb") as f:
        template_bytes = f.read()

    placeholders = _extract_placeholders(template_bytes)
    assert len(placeholders) > 200  # sanity: the rich template

    det_map, _audit = build_deterministic_fields(payload, job_id=payload.job_id)
    fallback_map = build_narrative_fallbacks(payload)
    field_map = _merge_like_handler(placeholders, det_map, fallback_map)

    filled = _fill_docx_template(template_bytes, field_map)

    # The core regression: not a single raw {{TOKEN}} may remain.
    assert _extract_placeholders(filled) == []

    # And the document actually carries real data, not just blanks.
    doc = Document(io.BytesIO(filled))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text
    assert payload.company_name in all_text


@pytest.mark.skipif(not os.path.exists(_SAMPLE), reason="fixture not present")
def test_every_deterministic_data_cell_has_a_value():
    """Deterministic builder must cover all non-narrative placeholders of the template."""
    payload = _load_payload()
    if not os.path.exists(_TEMPLATE):
        pytest.skip("template not present")
    with open(_TEMPLATE, "rb") as f:
        placeholders = _extract_placeholders(f.read())

    det_map, _audit = build_deterministic_fields(payload, job_id=payload.job_id)
    fallback_map = build_narrative_fallbacks(payload)

    missing = [
        ph for ph in placeholders
        if not is_narrative_field(ph) and ph not in det_map and ph not in fallback_map
    ]
    assert missing == [], f"deterministic placeholders not covered: {missing}"
