import io
import json
import logging
import os
import re
from collections import defaultdict
from datetime import datetime

import boto3

from shared.job_store import EXTRACTOR, REPORT_GENERATOR, RISK_SCORER, ANALYSIS_SUMMARY, save as job_save, save_bytes as job_save_bytes, load as job_load
from shared.agent_handoff import hydrate_input, slim_envelope
from shared.financial_math import is_cash_flow_account
from shared.llm_provider import LLMProvider
from shared.models import ScorerOutput, FinalReportOutput, NiifNoteDraft
from shared.niif_notes import load_niif_note
from shared.role_context import build_role_prompt_block, get_role_profile, get_role_label
from shared.s3_instructions import load_text
from shared.s3_report_store import (
    fetch_historical_reports,
    load_docx_template,
    slugify,
)
from .sheet_mapper import match_sheets
from .template_filler import (
    build_deterministic_fields,
    build_narrative_fallbacks,
    is_narrative_field,
)


logger = logging.getLogger("report_generator")
logger.setLevel(logging.INFO)

_RISK_ES = {"LOW": "bajo", "MEDIUM": "medio", "HIGH": "alto", "CRITICAL": "critico"}


def _money(value: float | int | None) -> str:
    if value is None:
        return "-"
    return f"{float(value):,.1f}"


def _pct(value: float | int | None) -> str:
    if value is None:
        return "-"
    sign = "+" if float(value) > 0 else ""
    return f"{sign}{float(value):.1f}%"


def _enum(value) -> str:
    return value.value if hasattr(value, "value") else str(value)


def _material_accounts(payload: ScorerOutput, limit: int = 8) -> list:
    # Exclude aggregate rows (grand_total / subtotal) so a "top finding" is a real
    # movement, not the sum of others; the same money would otherwise appear several
    # times (e.g. Total Activos, Total Activos Financieros and the portfolio line are
    # all the same ~1,540). Falls back to the full list for legacy outputs without
    # hierarchy tags. Breakdown_detail rows are kept — they are genuine positions.
    #
    # Also exclude movement rows (cash-flow / equity-changes statements): their
    # period-over-period % (contributions +238%, opening NAV +381%) is not an
    # interpretable position variation, so they must not appear in the "material
    # variations" list that drives the executive summary and LLM digest. Their
    # magnitude is surfaced separately through the fund-flow / cash-flow tables.
    candidates = [
        a for a in payload.analysis_results
        if getattr(a, "statement_role", "primary_line") not in ("grand_total", "subtotal")
        and not is_cash_flow_account(a)
    ] or list(payload.analysis_results)
    return sorted(
        candidates,
        key=lambda a: (
            0 if _enum(a.materiality) == "HIGH" else 1 if _enum(a.materiality) == "MEDIUM" else 2,
            -getattr(a, "impact_score", 0.0),
            -abs(a.variation_pct or 0),
        ),
    )[:limit]


def _build_executive_summary(payload: ScorerOutput) -> str:
    risk = _RISK_ES.get(_enum(payload.overall_risk_score), _enum(payload.overall_risk_score).lower())
    headline = (payload.risk_summary or {}).get("risk_headline")
    top = _material_accounts(payload, 3)
    top_text = "; ".join(f"{a.account_name}: {_pct(a.variation_pct)}" for a in top)
    base = payload.executive_narrative or headline or ""
    if base:
        return (
            f"{base} El perfil de riesgo global es {risk}, con score de validacion "
            f"{payload.validation_score}/100. Las variaciones mas relevantes son {top_text}."
        )
    return (
        f"{payload.company_name} presenta un perfil de riesgo {risk} y score de validacion "
        f"{payload.validation_score}/100. Las variaciones mas relevantes son {top_text}."
    )


def _build_board_summary(payload: ScorerOutput) -> str:
    totals = (payload.financial_ratios or {}).get("totals", {})
    ratios = (payload.financial_ratios or {}).get("ratios", {})
    risk_summary = payload.risk_summary or {}
    recs = risk_summary.get("risk_recommendations") or []
    rec_text = " ".join(recs[:3]) if recs else "Se recomienda mantener monitoreo periodico."

    # The board summary is the *detailed* counterpart to the punchier executive summary —
    # RevisorInteligente check 6.3 expects it to be richer. Lead with the risk headline, then
    # headline figures, a per-category risk breakdown, key ratios, and up to three
    # recommendations, so it is comprehensive rather than a two-line note.
    headline = (risk_summary.get("risk_headline") or "").strip()
    headline_text = f"{headline} " if headline else ""

    cat_text = ""
    cats = payload.risk_categories or {}
    if isinstance(cats, dict) and cats:
        parts = []
        for key, cat in cats.items():
            if not isinstance(cat, dict):
                continue
            level = _RISK_ES.get(_enum(cat.get("level")), str(cat.get("level") or "")).strip()
            score = cat.get("score")
            label = str(key).replace("_", " ").capitalize()
            score_txt = f" {score:.0f}/100" if isinstance(score, (int, float)) else ""
            parts.append(f"{label}: {level}{score_txt}".strip())
        if parts:
            cat_text = " Desglose de riesgo — " + "; ".join(parts) + "."

    ratio_bits = []
    if isinstance(ratios.get("roe_pct"), (int, float)):
        ratio_bits.append(f"ROE {ratios['roe_pct']:.1f}%")
    if isinstance(ratios.get("margen_neto_pct"), (int, float)):
        ratio_bits.append(f"margen neto {ratios['margen_neto_pct']:.1f}%")
    if isinstance(ratios.get("endeudamiento_global"), (int, float)):
        ratio_bits.append(f"endeudamiento {ratios['endeudamiento_global'] * 100:.1f}%")
    ratio_text = f" Indicadores clave: {', '.join(ratio_bits)}." if ratio_bits else ""

    return (
        f"{headline_text}"
        f"Activos: COP {_money(totals.get('total_assets'))} MM, "
        f"patrimonio: COP {_money(totals.get('total_equity'))} MM, "
        f"utilidad neta: COP {_money(totals.get('net_income'))} MM. "
        f"Riesgo {_RISK_ES.get(_enum(payload.overall_risk_score), _enum(payload.overall_risk_score).lower())}, "
        f"revision humana: {'si' if payload.requires_human_review else 'no'}."
        f"{cat_text}{ratio_text} {rec_text}"
    )


def _build_niif_notes(payload: ScorerOutput) -> list[NiifNoteDraft]:
    by_standard: dict[str, list] = defaultdict(list)
    for account in payload.analysis_results:
        if not account.requires_niif_note:
            continue
        refs = account.niif_note_references or payload.niif_notes_required or ["NIIF 7"]
        # Draft a note for EVERY standard the account references — not just the first
        # three. The previous `refs[:3]` cap dropped standards that sat past position 3
        # (e.g. "NIIF 7" on accounts that already cited NIC 1 / NIC 32 / NIIF 13 first),
        # yet the account kept the full list in `niif_note_references`. RevisorInteligente
        # check 3.1 then flagged every dropped standard as a cross-reference ERROR. Notes
        # are grouped by standard, so the count stays bounded by the ~10 distinct NIIF/NIC
        # standards regardless of how many refs any single account carries.
        for ref in refs:
            by_standard[ref].append(account)

    notes: list[NiifNoteDraft] = []
    for idx, (standard, accounts) in enumerate(sorted(by_standard.items()), start=1):
        top_accounts = accounts[:6]
        names = ", ".join(a.account_name for a in top_accounts)
        account_tail = (
            f" En este reporte la nota aplica a: {names}. "
            "Se basa exclusivamente en las cuentas marcadas por los agentes previos."
        )
        # Prefer the authored disclosure text for the standard (uploaded under
        # instructions/niif_notes/). Falls back to the generic template when no
        # note file is available, so a missing file never blocks the report.
        authored = load_niif_note(standard).strip()
        if authored:
            content = f"{authored}\n\n{account_tail.strip()}"
        else:
            content = (
                f"De acuerdo con {standard}, se recomienda revelar las variaciones materiales "
                f"asociadas a {names}. La medicion debe reconciliar saldos del periodo actual y "
                "comparativo, explicar cambios de valor razonable, liquidez o contraparte cuando "
                "aplique, y dejar evidencia de los juicios contables usados por la administracion. "
                "Esta nota se basa exclusivamente en las cuentas marcadas por los agentes previos."
            )
        notes.append(
            NiifNoteDraft(
                note_id=f"note-{idx:03d}",
                niif_reference=standard,
                title=f"Revelacion de variaciones materiales bajo {standard}",
                content=content,
                affected_account_ids=[a.account_id for a in top_accounts],
                requires_disclosure=True,
            )
        )
    return notes


# ---------------------------------------------------------------------------
# .docx template helpers
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
# Instruction markers ~~text~~ are stripped from the final document deterministically.
_INSTRUCTION_RE = re.compile(r"~~.+?~~", re.DOTALL)
# Detects row-based placeholder tables: PREFIX_Rn_Cm. The prefix may itself contain
# underscores (e.g. IS_ACC, FV_HIER, NAV_CLASS, NAV_FLOW); the `_R\d+_C\d+` suffix is
# fixed, so greedy matching resolves the prefix correctly via backtracking.
_ROW_PH_RE = re.compile(r"\{\{([A-Z0-9_]+)_R(\d+)_C\d+\}\}")
# Hard cap: never expand a table beyond this many data rows.
# Kept at 12 to stay within the 8-page target (sprint 4 polish).
_MAX_TABLE_ROWS = 12

_INCOME_KW = ("ingreso", "gasto", "costo", "utilidad", "resultado", "pérdida", "perdida", "comisi")


# ---------------------------------------------------------------------------
# Extractor output helpers
# ---------------------------------------------------------------------------

def _classify_extractor_account(acc: dict) -> str:
    """Classify a raw extractor account dict as balance_sheet or income_statement."""
    cat = (acc.get("category") or "").lower()
    if cat in ("revenue", "expense"):
        return "income_statement"
    name = (acc.get("normalized_account_name") or acc.get("raw_account_name") or "").lower()
    if any(k in name for k in _INCOME_KW):
        return "income_statement"
    return "balance_sheet"


def _compute_row_needs(
    extractor_accounts: list[dict], sheet_mapping: dict[str, str | None] | None = None
) -> dict[str, int]:
    """Return how many rows each dynamic table needs, per the matched Excel sheets.

    Counts come from the SAME sheet mapping that fills the tables
    (``get_accounts_for_canonical``), so every account in a mapped sheet gets a row.
    Each table's template default is kept as a floor. When no sheet mapping is
    available (PDF/CSV inputs without ``source_sheet``), falls back to name/keyword
    heuristics for the statement tables.
    """
    from .sheet_mapper import get_accounts_for_canonical
    from .template_filler import _classify_sheet

    # (canonical sheet key, table prefix, template-default rows, include totals)
    sheet_tables = [
        ("BALANCE", "BS", 6, True),
        ("ESTADO_RESULTADOS", "IS_ACC", 8, True),
        ("INSTRUMENTOS", "PORTFOLIO", 10, False),
        ("VALOR_RAZONABLE_ACTIVOS", "FV_HIER", 4, False),
        ("ACTIVOS_NETOS", "NAV_CLASS", 5, False),
        ("FLUJO_EFECTIVO", "NAV_FLOW", 4, False),
    ]
    needs: dict[str, int] = {}
    mapped_any = False
    for canonical, prefix, default, totals in sheet_tables:
        accs = get_accounts_for_canonical(
            extractor_accounts, canonical, sheet_mapping or {}, include_totals=totals
        )
        if accs:
            mapped_any = True
        needs[prefix] = max(default, len(accs))
    needs["IS_Q2"] = 4  # quarterly stub is always N/D — keep at template minimum

    if mapped_any:
        return needs

    # Fallback: no source_sheet at all — size the statement tables by heuristic.
    bs = [a for a in extractor_accounts
          if _classify_extractor_account(a) == "balance_sheet" and not a.get("is_total")]
    is_ = [a for a in extractor_accounts
           if _classify_extractor_account(a) == "income_statement" and not a.get("is_total")]
    inv = [a for a in extractor_accounts
           if a.get("investment_type")
           or "invers" in (a.get("normalized_account_name") or "").lower()]
    inv = [a for a in inv if not a.get("is_total")]
    needs["BS"] = max(6, len(bs))
    needs["IS_ACC"] = max(8, len(is_))
    needs["IS_Q2"] = 4  # quarterly stub is always N/D — keep at template minimum
    needs["PORTFOLIO"] = max(10, len(inv))
    return needs


def _load_extractor_accounts(job_id: str) -> list[dict]:
    """Load Agent 1 output from S3 and return the accounts list. Empty list on any error."""
    try:
        raw = job_load(job_id, EXTRACTOR)
        accounts = raw.get("accounts") or []
        logger.info("extractor_accounts | job=%s count=%d", job_id, len(accounts))
        return accounts
    except Exception as exc:
        logger.warning("extractor_accounts | job=%s load_failed=%s", job_id, exc)
        return []


# ---------------------------------------------------------------------------
# Dynamic table-row expansion
# ---------------------------------------------------------------------------

def _detect_table_row_info(table) -> tuple[str, int] | None:
    """Return (prefix, max_row_in_template) if the table uses PREFIX_Rn_Cm placeholders."""
    prefix: str | None = None
    max_row = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full = "".join(r.text for r in para.runs)
                for m in _ROW_PH_RE.finditer(full):
                    p, rn = m.group(1), int(m.group(2))
                    if prefix is None:
                        prefix = p
                    if rn > max_row:
                        max_row = rn
    return (prefix, max_row) if prefix and max_row > 0 else None


def _expand_docx_table(table, current_n: int, target_n: int, prefix: str) -> None:
    """Clone the last data row of *table* until the table has *target_n* data rows.

    The last row that contains ``{{PREFIX_R{current_n}_Cx}}`` is treated as the
    template row. Each clone renames those placeholders to the next row number
    so the field-map can fill every row deterministically.
    """
    import copy
    _W_T = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"

    # Find the tr element that anchors the last template row.
    anchor_tr = None
    for row in table.rows:
        texts = ["".join(r.text for r in para.runs)
                 for cell in row.cells for para in cell.paragraphs]
        if any(f"{prefix}_R{current_n}_" in t for t in texts):
            anchor_tr = row._tr
            break

    if anchor_tr is None:
        logger.warning("expand_table | prefix=%s anchor R%d not found", prefix, current_n)
        return

    old_marker = f"{prefix}_R{current_n}_"
    for new_r in range(current_n + 1, target_n + 1):
        new_tr = copy.deepcopy(anchor_tr)
        new_marker = f"{prefix}_R{new_r}_"
        for t_elem in new_tr.iter(_W_T):
            if t_elem.text and old_marker in t_elem.text:
                t_elem.text = t_elem.text.replace(old_marker, new_marker)
        table._tbl.append(new_tr)


def _expand_all_tables(docx_bytes: bytes, row_needs: dict[str, int]) -> bytes:
    """Return docx bytes with every PREFIX_Rn_Cm table expanded to meet *row_needs*."""
    try:
        from docx import Document as _DocxDoc
    except ImportError:
        return docx_bytes

    if not row_needs:
        return docx_bytes

    doc = _DocxDoc(io.BytesIO(docx_bytes))

    def _process(table):
        info = _detect_table_row_info(table)
        if info:
            prefix, current_n = info
            target_n = min(row_needs.get(prefix, 0), _MAX_TABLE_ROWS)
            if target_n > current_n:
                logger.info("expand_table | prefix=%s rows %d→%d", prefix, current_n, target_n)
                _expand_docx_table(table, current_n, target_n, prefix)
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    _process(nested)

    for table in doc.tables:
        _process(table)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_placeholders(docx_bytes: bytes) -> list[str]:
    """Return ordered unique placeholder names found in a .docx template."""
    try:
        from docx import Document as _DocxDoc
    except ImportError:
        logger.warning("python-docx not installed — docx template flow disabled")
        return []

    doc = _DocxDoc(io.BytesIO(docx_bytes))
    seen: list[str] = []
    seen_set: set[str] = set()

    def _scan_para(para):
        full_text = "".join(r.text for r in para.runs)
        for m in _PLACEHOLDER_RE.finditer(full_text):
            name = m.group(1)
            if name not in seen_set:
                seen_set.add(name)
                seen.append(name)

    def _scan_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_para(para)
                for nested in cell.tables:
                    _scan_table(nested)

    def _scan_paragraphs(paragraphs):
        for para in paragraphs:
            _scan_para(para)

    _scan_paragraphs(doc.paragraphs)
    for table in doc.tables:
        _scan_table(table)

    # headers and footers live in separate XML parts
    for section in doc.sections:
        for hf in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            if hf is None:
                continue
            _scan_paragraphs(hf.paragraphs)
            for table in hf.tables:
                _scan_table(table)

    return seen


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _render_into_paragraph(para, text: str) -> None:
    """Rewrite a paragraph's content with `text`, rendering Markdown **bold** spans
    and `\\n` as real Word line breaks, while preserving the paragraph's original
    run formatting (font, size, color taken from the template's first run)."""
    import copy
    from docx.oxml.ns import qn

    # Capture the template run's formatting so the rewritten text keeps the look.
    template_rpr = None
    if para.runs:
        rpr = para.runs[0]._element.find(qn("w:rPr"))
        if rpr is not None:
            template_rpr = copy.deepcopy(rpr)

    # Drop existing runs (paragraph-level style in pPr is retained).
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    def _styled_run():
        run = para.add_run()
        if template_rpr is not None:
            run._element.insert(0, copy.deepcopy(template_rpr))
        return run

    def _add(segment: str, *, bold: bool) -> None:
        # The `run.text` setter clears break/tab children, so a line break must be
        # its own run emitted before the text run — never combined with one.
        if not segment:
            return
        run = _styled_run()
        run.text = segment
        if bold:
            run.bold = True

    for line_idx, line in enumerate(text.split("\n")):
        if line_idx > 0:
            _styled_run().add_break()
        pos = 0
        for m in _BOLD_RE.finditer(line):
            if m.start() > pos:
                _add(line[pos:m.start()], bold=False)
            _add(m.group(1), bold=True)
            pos = m.end()
        _add(line[pos:], bold=False)


def _prune_blank_table_rows(doc) -> int:
    """Delete table rows whose every cell is blank after filling.

    Dynamic tables ship with a row floor (e.g. 8 income-statement rows) and the
    filler writes "" into rows it has no data for. Those empty rows render as ugly
    gaps, so they are removed here. A row is removed only when *all* its cells are
    empty — rows that carry a label or an "N/D" value (intentional content) stay.
    Header rows always carry static text and are never touched.
    """
    removed = 0

    def _process(table):
        nonlocal removed
        for row in list(table.rows):
            cells = row.cells
            if all(not c.text.strip() for c in cells):
                row._tr.getparent().remove(row._tr)
                removed += 1
                continue
            for cell in cells:
                for nested in cell.tables:
                    _process(nested)

    for table in doc.tables:
        _process(table)
    return removed


def _para_is_blank(p_elem) -> bool:
    """True if a <w:p> has no visible text and carries no image/object content.

    Page-break-only paragraphs are still "blank" by text, but they are handled
    separately by the caller so the break itself is never dropped.
    """
    from docx.oxml.ns import qn

    text = "".join(t.text or "" for t in p_elem.iter(qn("w:t")))
    if text.strip():
        return False
    for tag in ("w:drawing", "w:pict", "w:object"):
        if p_elem.find(".//" + qn(tag)) is not None:
            return False
    return True


def _para_has_page_break(p_elem) -> bool:
    """True if the paragraph forces a page break (pageBreakBefore or a <w:br type=page>)."""
    from docx.oxml.ns import qn

    ppr = p_elem.find(qn("w:pPr"))
    if ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None:
        return True
    for br in p_elem.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _prune_blank_paragraphs(doc) -> int:
    """Collapse runs of empty body paragraphs that produce blank pages / dead space.

    The template (and the placeholder fill, which clears a paragraph to "" when a
    narrative field is absent) leaves many empty `<w:p>` paragraphs. Stacked, they
    render as large vertical gaps and even fully blank pages.

    Rules (body-level paragraphs only — table rows are handled by
    `_prune_blank_table_rows`):
      * At most ONE empty spacer paragraph is kept between two pieces of content.
      * Page-break paragraphs are always preserved (they drive section layout) and
        do not count against the single-spacer allowance.
      * Leading empty paragraphs (before any content) and trailing empty spacers
        (after the last content, before the final sectPr) are removed entirely.
      * Tables reset the spacer allowance, so one spacer after a table is allowed.
    """
    body = doc.element.body
    children = list(body.iterchildren())
    removed = 0
    last_was_blank_spacer = True  # start True → strip leading blanks

    for child in children:
        tag = child.tag.split("}")[-1]
        if tag != "p":
            # Tables / sectPr reset the run; a single spacer after them is allowed.
            last_was_blank_spacer = False
            continue
        if not _para_is_blank(child):
            last_was_blank_spacer = False
            continue
        # Blank paragraph.
        if _para_has_page_break(child):
            # Keep the break; treat it as a spacer so adjacent blanks collapse.
            last_was_blank_spacer = True
            continue
        if last_was_blank_spacer:
            child.getparent().remove(child)
            removed += 1
        else:
            last_was_blank_spacer = True

    # Strip any trailing blank spacers left before the final sectPr.
    for child in reversed(list(body.iterchildren())):
        tag = child.tag.split("}")[-1]
        if tag == "sectPr":
            continue
        if tag == "p" and _para_is_blank(child) and not _para_has_page_break(child):
            child.getparent().remove(child)
            removed += 1
            continue
        break

    return removed


def _fill_docx_template(docx_bytes: bytes, field_map: dict[str, str]) -> bytes:
    """Replace {{PLACEHOLDER}} markers in a .docx with the provided field values."""
    try:
        from docx import Document as _DocxDoc
    except ImportError:
        return docx_bytes

    doc = _DocxDoc(io.BytesIO(docx_bytes))

    def _replace_paragraph(para):
        full_text = "".join(r.text for r in para.runs)
        has_placeholder = "{{" in full_text
        has_instruction = "~~" in full_text
        if not has_placeholder and not has_instruction:
            return
        modified = full_text
        if has_placeholder:
            modified = _PLACEHOLDER_RE.sub(lambda m: field_map.get(m.group(1), "") or "", modified)
        if has_instruction:
            modified = _INSTRUCTION_RE.sub("", modified)
        if modified != full_text:
            if not modified.strip():
                for run in list(para.runs):
                    run._element.getparent().remove(run._element)
            else:
                _render_into_paragraph(para, modified)

    def _replace_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_paragraph(para)
                for nested in cell.tables:
                    _replace_table(nested)

    for para in doc.paragraphs:
        _replace_paragraph(para)
    for table in doc.tables:
        _replace_table(table)

    # headers and footers live in separate XML parts — iterate all sections
    for section in doc.sections:
        for hf in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            if hf is None:
                continue
            for para in hf.paragraphs:
                _replace_paragraph(para)
            for table in hf.tables:
                _replace_table(table)

    pruned = _prune_blank_table_rows(doc)
    if pruned:
        logger.info("prune_blank_rows | removed=%d", pruned)

    pruned_paras = _prune_blank_paragraphs(doc)
    if pruned_paras:
        logger.info("prune_blank_paragraphs | removed=%d", pruned_paras)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# System prompt
# ---------------------------------------------------------------------------

_PROMPT_S3_KEY = "instructions/prompts/04_prompt_agent_report-generator.md"
_LOCAL_PROMPT_PATH = os.path.normpath(
    os.path.join(os.path.dirname(__file__), "..", "system_pompts",
                 "04_prompt_agent_report-generator.md")
)
_INLINE_FALLBACK = (
    "Eres el agente generador de reportes financieros institucionales de CreditIQ. "
    "Recibes una lista de campos de plantilla y datos financieros. "
    "Responde SOLO con un objeto JSON donde cada clave es el nombre del campo "
    "y el valor es el contenido institucional a insertar en ese campo."
)
_prompt_cache: str | None = None


def _get_system_prompt() -> str:
    global _prompt_cache
    if _prompt_cache is not None:
        return _prompt_cache
    s3_text = load_text(_PROMPT_S3_KEY, fallback="")
    if s3_text:
        _prompt_cache = s3_text
        logger.info("report_prompt | source=s3 chars=%d", len(s3_text))
        return _prompt_cache
    try:
        with open(_LOCAL_PROMPT_PATH, encoding="utf-8") as f:
            local_text = f.read()
        if len(local_text) > 100:
            _prompt_cache = local_text
            logger.info("report_prompt | source=local chars=%d", len(local_text))
            return _prompt_cache
    except OSError:
        pass
    logger.warning("report_prompt | source=inline_fallback")
    _prompt_cache = _INLINE_FALLBACK
    return _prompt_cache


# ---------------------------------------------------------------------------
# LLM field generation
# ---------------------------------------------------------------------------

def _build_llm_digest(payload: ScorerOutput, historical: list, extractor_accounts: list[dict] | None = None) -> str:
    """Build a compact text digest for the LLM (~3000 tokens max)."""
    risk = _enum(payload.overall_risk_score)
    health = _enum(payload.overall_financial_health)
    periods_str = " / ".join(payload.periods) if payload.periods else "-"
    conf_label = "Alta" if payload.analysis_confidence >= 0.8 else "Media" if payload.analysis_confidence >= 0.6 else "Baja"

    lines = [
        f"# BRIEFING: {payload.company_name}",
        f"Período: {periods_str} | Moneda: {payload.currency}",
        f"Riesgo Global: {risk} | Salud Financiera: {health}",
        f"Score Validación: {payload.validation_score}/100 | Confianza: {conf_label}",
        f"Anti-Alucinación: {'PASSED' if payload.anti_hallucination_passed else 'FAILED'}",
        "",
    ]

    kpis = payload.executive_kpis or {}
    if kpis:
        lines.append("## KPIs EJECUTIVOS")
        for k, v in list(kpis.items())[:8]:
            lines.append(f"- {k}: {v}")
        lines.append("")

    top = _material_accounts(payload, 10)
    if top:
        lines.append("## CUENTAS MATERIALES (TOP 10)")
        lines.append("| Cuenta | Actual | Anterior | Δ% | Materialidad | Señal |")
        lines.append("|--------|-------:|---------:|----:|:------------:|-------|")
        for a in top:
            signal = getattr(a, "investment_signal", None) or ""
            anomaly = " ⚠" if a.anomaly_detected else ""
            lines.append(
                f"| {a.account_name}{anomaly} | {_money(a.current_value)} "
                f"| {_money(a.previous_value)} | {_pct(a.variation_pct)} "
                f"| {_enum(a.materiality)} | {signal} |"
            )
        lines.append("")

    categories = payload.risk_categories or {}
    narratives = (payload.risk_summary or {}).get("category_narratives") or {}
    if categories:
        lines.append("## PERFIL DE RIESGO")
        for key in ("credito", "mercado", "financiero"):
            cat = categories.get(key)
            if not cat:
                continue
            label = cat.get("label", key)
            level = cat.get("level", "N/A")
            score = cat.get("score", 0)
            narrative = narratives.get(key, "")
            findings = cat.get("key_findings") or []
            lines.append(f"### {label}: {level} ({score}/100)")
            if narrative:
                lines.append(narrative)
            for f in findings[:3]:
                lines.append(f"- {f}")
        lines.append("")

    synthesis = payload.executive_synthesis or {}
    if payload.portfolio_thesis:
        lines.append("## TESIS DE PORTAFOLIO")
        lines.append(payload.portfolio_thesis)
        lines.append("")
    if synthesis.get("portfolio_story"):
        lines.append("## HISTORIA DEL PORTAFOLIO")
        lines.append(synthesis["portfolio_story"])
        lines.append("")
    board_alerts = synthesis.get("board_alerts") or []
    if board_alerts:
        lines.append("## ALERTAS JUNTA")
        for alert in board_alerts[:5]:
            lines.append(f"- {alert}")
        lines.append("")

    sc = payload.sheet_concentration or {}
    if sc.get("asset_available") or sc.get("instrument_available") or sc.get("bank_available"):
        lines.append("## COMPOSICIÓN DEL PORTAFOLIO")
        if sc.get("asset_available"):
            asset_bd = (sc.get("asset_breakdown") or [])[:5]
            lines.append("Activos: " + " | ".join(f"{r.get('name')}: {_pct(r.get('pct'))}" for r in asset_bd))
        if sc.get("instrument_available"):
            inst_bd = (sc.get("instrument_breakdown") or [])[:5]
            lines.append("Instrumentos: " + " | ".join(f"{r.get('instrument_type')}: {_pct(r.get('pct'))}" for r in inst_bd))
        if sc.get("bank_available"):
            bank_bd = (sc.get("bank_breakdown") or [])[:5]
            lines.append("Bancos/Custodios: " + " | ".join(f"{r.get('name')}: {_pct(r.get('pct'))}" for r in bank_bd))
        lines.append("")

    signals = payload.cross_statement_signals or []
    if signals:
        lines.append("## SEÑALES CROSS-STATEMENT")
        for s in signals[:5]:
            desc = s.get("description") or s.get("finding") or ""
            implication = s.get("implication") or ""
            if desc:
                lines.append(f"- {desc} {implication}".strip())
        lines.append("")

    anomalies = [a for a in payload.analysis_results if a.anomaly_detected]
    if anomalies:
        lines.append("## ANOMALÍAS DETECTADAS")
        for a in anomalies[:8]:
            insight = a.executive_insight or ""
            lines.append(f"- **{a.account_name}** ({_pct(a.variation_pct)}): {insight}")
        lines.append("")

    niif_refs = payload.niif_notes_required or []
    if niif_refs:
        lines.append("## REQUERIMIENTOS NIIF")
        lines.append(", ".join(niif_refs[:10]))
        lines.append("")

    macro = payload.macro_context or {}
    if macro:
        lines.append("## CONTEXTO MACROECONÓMICO")
        for k, v in list(macro.items())[:8]:
            lines.append(f"- {k}: {v}")
        lines.append("")

    if historical:
        lines.append("## CONTEXTO HISTÓRICO")
        lines.append(f"Se dispone de {len(historical)} reporte(s) histórico(s) comparables.")
        lines.append("")

    risk_summary = payload.risk_summary or {}
    if risk_summary.get("risk_headline"):
        lines.append("## TITULAR DE RIESGO")
        lines.append(risk_summary["risk_headline"])
        lines.append("")
    recs = risk_summary.get("risk_recommendations") or []
    if recs:
        lines.append("## RECOMENDACIONES")
        for r in recs[:5]:
            lines.append(f"- {r}")
        lines.append("")

    fpa = payload.fund_policy_assessment or {}
    if fpa and fpa.get("assessments"):
        lines.append("## POLÍTICA DE FONDO (LÍMITES REGULATORIOS)")
        lines.append(f"Tipo de fondo: {fpa.get('fund_type', 'N/A')} | Fuente: {fpa.get('policy_source', 'default')}")
        lines.append(f"Resumen: {fpa.get('summary', '')}")
        limits = fpa.get("limits", {})
        if limits:
            lines.append(
                f"Límites: por emisor {limits.get('per_issuer_pct', '?')}% | "
                f"sector {limits.get('per_sector_pct', '?')}% | "
                f"contraparte {limits.get('per_counterparty_pct', '?')}%"
            )
        breaches = [a for a in fpa.get("assessments", []) if a.get("status") == "breach"]
        nears = [a for a in fpa.get("assessments", []) if a.get("status") == "near"]
        if breaches:
            lines.append("INCUMPLIMIENTOS:")
            for b in breaches[:5]:
                lines.append(
                    f"  - {b.get('dimension')} / {b.get('name')}: "
                    f"{b.get('exposure_pct', 0):.1f}% vs límite {b.get('limit_pct', 0):.1f}% "
                    f"(exceso: {abs(b.get('headroom_pct', 0)):.1f}%)"
                )
        if nears:
            lines.append("CERCA DEL LÍMITE:")
            for n in nears[:5]:
                lines.append(
                    f"  - {n.get('dimension')} / {n.get('name')}: "
                    f"{n.get('exposure_pct', 0):.1f}% vs límite {n.get('limit_pct', 0):.1f}% "
                    f"(margen: {n.get('headroom_pct', 0):.1f}%)"
                )
        if not breaches and not nears:
            lines.append("Todas las dimensiones dentro de los límites permitidos.")
        lines.append("INSTRUCCIÓN: Menciona siempre el porcentaje de exposición y el límite permitido al redactar concentraciones.")
        lines.append("")

    tv = payload.top_variations or []
    if tv:
        lines.append("## TOP VARIACIONES CON PERÍODO COMPARATIVO")
        lines.append("| Cuenta | Período actual | Período comparativo | Δ Absoluta | Δ% | Materialidad |")
        lines.append("|--------|:--------------:|:-------------------:|----------:|----:|:------------:|")
        for v in tv[:10]:
            cur_p = v.get("current_period", "?")
            cmp_p = v.get("comparative_period", "?")
            lines.append(
                f"| {v.get('account_name', '')} | {cur_p} | {cmp_p} "
                f"| {_money(v.get('absolute_variation', 0))} "
                f"| {_pct(v.get('variation_pct'))} | {v.get('materiality', 'LOW')} |"
            )
        lines.append("INSTRUCCIÓN: Al nombrar variaciones usa siempre 'variación entre {período_actual} y {período_comparativo}'.")
        lines.append("")

    cb = payload.comparative_basis or {}
    if cb:
        lines.append("## BASES COMPARATIVAS (IFRS)")
        lines.append("REGLA: Al mencionar cualquier variación debes nombrar ambos períodos explícitamente.")
        for stmt, entry in cb.items():
            cur = entry.get("current", "?")
            comp = entry.get("comparative", "?")
            rule = entry.get("rule", "?")
            warn = entry.get("mismatch_warning", "")
            lines.append(f"- {stmt}: {cur} vs {comp} [{rule}]" + (f" ⚠ {warn}" if warn else ""))
        lines.append("")

    # Extractor accounts grouped by source_sheet — gives the LLM the raw sheet
    # structure that the Word template tables directly mirror (e.g. Inversiones,
    # Balance, Resultados). Capped per sheet to stay within token budget.
    if extractor_accounts:
        from collections import defaultdict as _dd
        by_sheet: dict[str, list] = _dd(list)
        for acc in extractor_accounts:
            sheet = acc.get("source_sheet") or "Sin hoja"
            by_sheet[sheet].append(acc)
        lines.append("## CUENTAS POR HOJA (EXTRACTOR)")
        for sheet, accs in sorted(by_sheet.items()):
            lines.append(f"### Hoja: {sheet} ({len(accs)} cuentas)")
            for acc in accs[:20]:  # cap per sheet
                cur = acc.get("current_value")
                prev = acc.get("previous_value")
                name = acc.get("normalized_account_name") or acc.get("raw_account_name") or "—"
                inv_type = acc.get("investment_type") or ""
                issuer = acc.get("issuer_name") or ""
                label = f"{name}"
                if issuer:
                    label += f" ({issuer})"
                if inv_type:
                    label += f" [{inv_type}]"
                cur_s = f"{float(cur):,.1f}" if cur is not None else "N/D"
                prev_s = f"{float(prev):,.1f}" if prev is not None else "N/D"
                lines.append(f"- {label}: {cur_s} / {prev_s}")
            if len(accs) > 20:
                lines.append(f"  … y {len(accs) - 20} más")
        lines.append("")

    return "\n".join(lines)


def _call_llm_for_fields(
    payload: ScorerOutput,
    digest: str,
    placeholders: list[str],
    llm_model: str | None = None,
) -> dict[str, str]:
    """Ask the LLM to fill each template placeholder. Returns {} on error."""
    system_prompt = _get_system_prompt()
    field_list = "\n".join(f"- {{{{{p}}}}}" for p in placeholders)

    # AI Analysis Perspectives — tell the narrative author which sections deserve
    # more (or less) development for the selected professional role.
    emphasis_note = ""
    profile = get_role_profile(payload.analysis_role)
    emphasis = profile.get("report_emphasis") or {}
    if payload.analysis_role != "general" and (emphasis.get("boost") or emphasis.get("reduce")):
        boost = ", ".join(emphasis.get("boost", []))
        reduce = ", ".join(emphasis.get("reduce", []))
        emphasis_note = (
            f"## ÉNFASIS DEL REPORTE (perfil: {profile.get('label', payload.analysis_role)})\n\n"
            f"Desarrolla con mayor profundidad: {boost or 'N/A'}.\n"
            f"Trata con menor detalle (sin omitir hallazgos críticos): {reduce or 'N/A'}.\n\n"
        )

    user_prompt = (
        f"{emphasis_note}"
        f"## CAMPOS DEL TEMPLATE\n\n{field_list}\n\n"
        f"## DATOS FINANCIEROS\n\n{digest}"
    )
    llm = LLMProvider(model=llm_model)
    llm.set_role_context(build_role_prompt_block(payload.analysis_role))
    try:
        result = llm.generate_json(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            temperature=0.5,
            tenant_id=payload.tenant_id,
            job_id=payload.job_id,
            max_tokens=12000,
        )
        valid = {k: str(v) for k, v in result.items() if k in set(placeholders)}
        logger.info(
            "llm_for_fields | job=%s placeholders=%d filled=%d",
            payload.job_id, len(placeholders), len(valid),
        )
        return valid
    except Exception as exc:
        logger.error("llm_for_fields failed | job=%s error=%s", payload.job_id, exc)
        return {}


def _deduplicate_narrative_fields(llm_map: dict[str, str]) -> dict[str, str]:
    """Remove sentences from executive/conclusion fields that already appear verbatim
    in a finding body. Findings are the most specific section and win on conflict.

    Priority (highest keeps the sentence):
      1. FINDING_*_BODY
      2. NEXT_STEPS / CONCLUSION fields
      3. EXEC_CONCLUSIONS / RESUMEN fields (pruned last)
    """
    def _sentences(text: str) -> list[str]:
        """Split text into non-trivial sentences (≥ 6 words)."""
        parts = re.split(r"(?<=[.!?])\s+", (text or "").strip())
        return [s.strip() for s in parts if len(s.split()) >= 6]

    def _normalize(s: str) -> str:
        return re.sub(r"\s+", " ", s.lower().strip().rstrip(".!?"))

    # Collect all sentence fingerprints from the highest-priority sections.
    seen: set[str] = set()
    for key, val in llm_map.items():
        if re.search(r"FINDING_\d+_BODY", key, re.IGNORECASE):
            for s in _sentences(val):
                seen.add(_normalize(s))

    if not seen:
        return llm_map  # nothing to deduplicate against

    result = dict(llm_map)
    # Fields that should be pruned if they repeat finding sentences
    _EXEC_KEYS = re.compile(r"EXEC|RESUMEN|SINTESIS|EJECUTIVO|SUMMARY|NEXT_STEP|CONCLUSION|CIERRE|PERSPECTIVA", re.IGNORECASE)
    for key, val in llm_map.items():
        if not _EXEC_KEYS.search(key):
            continue
        sentences = _sentences(val)
        kept = [s for s in sentences if _normalize(s) not in seen]
        if len(kept) < len(sentences):
            result[key] = " ".join(kept) if kept else val  # preserve original if all removed
            logger.debug("dedup | field=%s removed=%d/%d sentences", key, len(sentences) - len(kept), len(sentences))
    return result


# ---------------------------------------------------------------------------
# Lambda entry point
# ---------------------------------------------------------------------------

_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_analysis_summary_md(payload: ScorerOutput, executive_summary: str, board_summary: str) -> str:
    """Build a self-contained markdown summary of the full analysis.

    Saved as analysis_summary.md alongside the other job artifacts so it can be
    loaded cheaply on any future request (audit, chat, re-review) without having
    to re-fetch and parse all agent outputs.
    """
    def _fmt(v) -> str:
        if v is None:
            return "N/D"
        try:
            return f"{float(v):,.1f}"
        except (TypeError, ValueError):
            return str(v)

    def _pct_fmt(v) -> str:
        if v is None:
            return "N/D"
        try:
            f = float(v)
            return f"{'+' if f > 0 else ''}{f:.1f}%"
        except (TypeError, ValueError):
            return str(v)

    periods = " / ".join(payload.periods) if payload.periods else "N/D"
    totals = (payload.financial_ratios or {}).get("totals") or {}
    ratios = (payload.financial_ratios or {}).get("ratios") or {}
    rs = payload.risk_summary or {}
    cats = payload.risk_categories or {}
    synth = payload.executive_synthesis or {}
    fa = payload.fund_analysis or {}
    kpis = payload.executive_kpis or {}

    lines: list[str] = [
        f"# Resumen de Análisis — {payload.company_name}",
        f"**Períodos:** {periods}  |  **Moneda:** {payload.currency or 'COP'}  "
        f"|  **Generado:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        f"**Job ID:** {payload.job_id}",
    ]
    if payload.analysis_role != "general":
        lines.append(f"**Perspectiva de análisis:** {get_role_label(payload.analysis_role)}")
    lines += [
        "",
        "---",
        "",
        "## Resumen Ejecutivo",
        executive_summary,
        "",
        "## Resumen para Junta Directiva",
        board_summary,
        "",
        "---",
        "",
        "## Perfil de Riesgo",
        f"- **Riesgo global:** {payload.overall_risk_score}",
        f"- **Salud financiera:** {payload.overall_financial_health}",
        f"- **Score de validación:** {payload.validation_score}/100",
        f"- **Requiere revisión humana:** {'Sí' if payload.requires_human_review else 'No'}",
    ]

    for key, cat in cats.items():
        findings = "; ".join((cat.get("key_findings") or [])[:2])
        lines.append(f"- **{cat.get('label', key)}:** {cat.get('level')} ({cat.get('score', 0)}/100) — {findings}")

    if rs.get("risk_headline"):
        lines += ["", f"**Titular de riesgo:** {rs['risk_headline']}"]

    recs = rs.get("risk_recommendations") or []
    if recs:
        lines += ["", "**Recomendaciones:**"]
        for r in recs[:5]:
            lines.append(f"- {r}")

    lines += ["", "---", "", "## KPIs Ejecutivos"]
    for k, v in list(kpis.items())[:10]:
        lines.append(f"- **{k}:** {v}")

    lines += ["", "---", "", "## Cifras Clave (COP miles)"]
    for label, key in [
        ("Total Activos", "total_assets"),
        ("Total Pasivos", "total_liabilities"),
        ("Patrimonio / Activos Netos", "total_equity"),
        ("Ingresos / Resultado", "total_revenue"),
        ("Utilidad Neta", "net_income"),
    ]:
        lines.append(f"- **{label}:** {_fmt(totals.get(key))}")

    for label, key in [("ROE", "roe_pct"), ("ROA", "roa_pct"), ("Deuda/Patrimonio", "debt_to_equity")]:
        v = ratios.get(key)
        if v is not None:
            lines.append(f"- **{label}:** {_pct_fmt(v)}")

    if fa.get("is_fund"):
        recon = fa.get("nav_reconciliation") or {}
        lines += [
            "", "**Fondo de inversión — NAV:**",
            f"- NAV apertura: {_fmt(recon.get('opening_nav'))}",
            f"- NAV cierre: {_fmt(recon.get('closing_nav'))}",
            f"- Flujo neto: {_fmt(recon.get('net_investor_flow'))}",
        ]

    # Material accounts. AccountAnalysis has no `is_total` — that field lives on the
    # raw ExtractedAccount. Exclude aggregate rows via the hierarchy tag instead
    # (same approach as _material_accounts); referencing a.is_total raised an
    # AttributeError that silently dropped analysis_summary.md on every run.
    material = sorted(
        [a for a in payload.analysis_results
         if getattr(a, "statement_role", "primary_line") not in ("grand_total", "subtotal")
         and not is_cash_flow_account(a)],
        key=lambda a: (0 if _enum(a.materiality) == "HIGH" else 1, -(a.impact_score or 0)),
    )[:15]
    if material:
        lines += ["", "---", "", "## Cuentas Materiales", "| Cuenta | Actual | Anterior | Δ% | Materialidad |",
                  "|--------|-------:|---------:|----:|:------------:|"]
        for a in material:
            lines.append(
                f"| {a.account_name} | {_fmt(a.current_value)} | {_fmt(a.previous_value)} "
                f"| {_pct_fmt(a.variation_pct)} | {_enum(a.materiality)} |"
            )

    # Anomalies
    anomalies = [a for a in payload.analysis_results if a.anomaly_detected][:8]
    if anomalies:
        lines += ["", "---", "", "## Anomalías Detectadas"]
        for a in anomalies:
            insight = a.executive_insight or ""
            lines.append(f"- **{a.account_name}** (Δ{_pct_fmt(a.variation_pct)}): {insight}")

    # Portfolio story
    if synth.get("portfolio_story"):
        lines += ["", "---", "", "## Historia del Portafolio", synth["portfolio_story"]]

    if payload.portfolio_thesis:
        lines += ["", "**Tesis de portafolio:** " + payload.portfolio_thesis]

    board_alerts = synth.get("board_alerts") or []
    if board_alerts:
        lines += ["", "**Alertas para la Junta:**"]
        for alert in board_alerts:
            lines.append(f"- {alert}")

    # NIIF
    niif_refs = payload.niif_notes_required or []
    if niif_refs:
        lines += ["", "---", "", f"## Requerimientos NIIF", ", ".join(niif_refs)]

    lines += ["", "---", f"*Generado por CreditIQ Multi-Agente · Job {payload.job_id}*"]
    return "\n".join(lines)




def lambda_handler(event: dict, context) -> dict:
    """
    Input:  ScorerOutput
    Output: FinalReportOutput — primary deliverable is a filled .docx stored in the
            job S3 folder at jobs/{date}/{job_id}/report_generator_report.docx

    Flow:
      1. Load CreditIQ_Template_EEFF.docx from S3
      2. Extract {{PLACEHOLDER}} fields from the template
      3. Call LLM → JSON mapping of field names to institutional content
      4. Fill template bytes with the field map
      5. Save filled .docx to jobs/{date}/{job_id}/report_generator_report.docx
      6. Store s3://bucket/key in FinalReportOutput.docx_report_url
    """
    # llm_model is a transient field carried on the event (not in the S3 artifact);
    # read it before rehydrating the full ScorerOutput from S3. See agent_handoff.py.
    llm_model: str | None = event.get("llm_model") or None
    payload = hydrate_input(event, RISK_SCORER, ScorerOutput, discriminator="analysis_results")
    bucket = os.environ["MAIN_BUCKET"]
    reference_date = datetime.utcnow()

    logger.info("ReportGenerator start | job=%s company=%s", payload.job_id, payload.company_name)

    historical = fetch_historical_reports(
        tenant_id=payload.tenant_id,
        company_slug=slugify(payload.company_name),
        reference_date=reference_date,
        bucket=bucket,
    )

    # Load Agent 1 raw accounts to drive dynamic table expansion and richer LLM context.
    extractor_accounts = _load_extractor_accounts(payload.job_id)

    # Resolve actual Excel sheet names → canonical template table keys.
    # The LLM is used only for sheet names that keyword matching cannot classify.
    actual_sheets = sorted({
        acc["source_sheet"]
        for acc in extractor_accounts
        if acc.get("source_sheet")
    })
    sheet_mapping = match_sheets(
        actual_sheets,
        llm=LLMProvider(model=llm_model),
        tenant_id=payload.tenant_id,
        job_id=payload.job_id,
    ) if actual_sheets else {}

    # Row counts for dynamic table expansion are derived from the SAME sheet mapping
    # that fills the tables, so every account in a mapped sheet gets a row.
    row_needs = _compute_row_needs(extractor_accounts, sheet_mapping) if extractor_accounts else {}

    digest = _build_llm_digest(payload, historical, extractor_accounts=extractor_accounts)

    # ── .docx template flow ──────────────────────────────────────────────────
    docx_s3_url: str | None = None
    template_bytes = load_docx_template(bucket)

    if template_bytes:
        # Expand tables that have fewer rows than the extractor found accounts for.
        # This must happen BEFORE placeholder extraction so the new rows are discovered.
        if row_needs:
            template_bytes = _expand_all_tables(template_bytes, row_needs)

        placeholders = _extract_placeholders(template_bytes)
        narrative = [p for p in placeholders if is_narrative_field(p)]
        logger.info(
            "docx_template | job=%s placeholders=%d narrative=%d row_needs=%s",
            payload.job_id, len(placeholders), len(narrative), row_needs,
        )

        # 1. Deterministic calculated cells — the bulk of the template (tables, KPIs).
        det_map, field_audit = build_deterministic_fields(
            payload, generated_at=reference_date, job_id=payload.job_id,
            row_needs=row_needs, extractor_accounts=extractor_accounts,
            sheet_mapping=sheet_mapping,
        )
        # 2. Deterministic narrative fallbacks — used if the LLM omits/fails a field.
        fallback_map = build_narrative_fallbacks(payload)
        # 3. LLM authors ONLY the narrative sections (never the data cells).
        llm_map = (
            _call_llm_for_fields(payload, digest, narrative, llm_model=llm_model)
            if narrative else {}
        )

        # 4. De-duplication pass: ensure each conclusion appears in exactly one section.
        #    Findings are most specific (kept); executive summary and next_steps lose
        #    sentences already present verbatim in a finding body.
        llm_map = _deduplicate_narrative_fields(llm_map)

        # Merge per placeholder. Precedence: LLM > deterministic > fallback > "".
        # Every extracted placeholder gets a value, so no raw {{TOKEN}} leaks through
        # and a .docx is always produced once the template loads.
        field_map: dict[str, str] = {}
        for ph in placeholders:
            llm_val = llm_map.get(ph)
            if llm_val and llm_val.strip():
                field_map[ph] = llm_val
                field_audit[ph] = "llm_narrative"
            elif ph in det_map:
                field_map[ph] = det_map[ph]
            elif ph in fallback_map:
                field_map[ph] = fallback_map[ph]
                field_audit[ph] = "deterministic_fallback"
            else:
                field_map[ph] = ""
                field_audit[ph] = "N/D (no source)"

        # Build fill summary for the audit log
        nd_fields = [k for k, v in field_map.items() if not v or v == "N/D"]
        filled_fields = [k for k, v in field_map.items() if v and v != "N/D"]
        audit_payload = {
            "job_id": payload.job_id,
            "generated_at": reference_date.isoformat(),
            "sheet_mapping": {k: v for k, v in (sheet_mapping or {}).items()},
            "summary": {
                "total_placeholders": len(placeholders),
                "filled": len(filled_fields),
                "nd_gaps": len(nd_fields),
                "llm_narrative_filled": sum(1 for p in narrative if (llm_map.get(p) or "").strip()),
            },
            "nd_fields": nd_fields,
            "provenance": field_audit,
        }
        try:
            job_save(payload.job_id, "report_generator_audit", audit_payload)
            logger.info(
                "audit_saved | job=%s filled=%d nd=%d",
                payload.job_id, len(filled_fields), len(nd_fields),
            )
        except Exception as exc:
            logger.warning("audit_save_failed | job=%s error=%s", payload.job_id, exc)

        filled_bytes = _fill_docx_template(template_bytes, field_map)
        docx_key = job_save_bytes(
            payload.job_id,
            "report_generator_report",
            filled_bytes,
            _DOCX_CONTENT_TYPE,
            "docx",
        )
        docx_s3_url = f"s3://{bucket}/{docx_key}"
        llm_filled = audit_payload["summary"]["llm_narrative_filled"]
        logger.info(
            "docx_saved | job=%s key=%s filled=%d nd=%d llm_narrative=%d/%d",
            payload.job_id, docx_key, len(filled_fields), len(nd_fields),
            llm_filled, len(narrative),
        )
    else:
        logger.warning("docx_template_not_found | job=%s bucket=%s", payload.job_id, bucket)

    # ── Build FinalReportOutput (flows to RevisorInteligente) ────────────────
    result = FinalReportOutput(
        job_id=payload.job_id,
        tenant_id=payload.tenant_id,
        analysis_role=payload.analysis_role,
        company_name=payload.company_name,
        periods=payload.periods,
        generated_at=reference_date,
        validation_score=payload.validation_score,
        overall_risk_score=payload.overall_risk_score,
        overall_financial_health=payload.overall_financial_health,
        executive_summary=_build_executive_summary(payload),
        board_summary=_build_board_summary(payload),
        analysis_results=payload.analysis_results,
        niif_note_drafts=_build_niif_notes(payload),
        docx_report_url=docx_s3_url,
        risk_categories=payload.risk_categories,
        risk_summary=payload.risk_summary,
        financial_ratios=payload.financial_ratios,
        fund_analysis=payload.fund_analysis,
        executive_kpis=payload.executive_kpis,
        sheet_concentration=payload.sheet_concentration,
        structured_analysis=payload.structured_analysis,
        cross_statement_signals=payload.cross_statement_signals,
        earnings_sustainability=payload.earnings_sustainability,
        comparative_basis=payload.comparative_basis,
        fund_policy_assessment=payload.fund_policy_assessment,
        top_variations=payload.top_variations,
    )

    # ── Save analysis_summary.md ─────────────────────────────────────────────
    # A self-contained markdown summary of the full analysis, saved alongside the
    # other job artifacts. Used by the chat handler and audit workflows to load
    # rich context cheaply without re-fetching all agent outputs.
    try:
        summary_md = _build_analysis_summary_md(payload, result.executive_summary, result.board_summary)
        job_save_bytes(result.job_id, ANALYSIS_SUMMARY, summary_md.encode("utf-8"), "text/markdown", "md")
        logger.info("analysis_summary_saved | job=%s chars=%d", result.job_id, len(summary_md))
    except Exception as exc:
        logger.warning("analysis_summary_save_failed | job=%s error=%s", result.job_id, exc)

    output = result.model_dump(mode="json")
    try:
        job_save(result.job_id, REPORT_GENERATOR, output)
    except Exception as exc:
        logger.error("Failed to persist report generator output: %s", exc)

    logger.info(
        "ReportGenerator done | job=%s docx=%s",
        payload.job_id, docx_s3_url or "none",
    )
    # Full FinalReportOutput is persisted to S3 above; return only a claim-check
    # pointer so the SFN handoff to RevisorInteligente stays under the 256 KB cap.
    return slim_envelope(result, status="report_complete")
